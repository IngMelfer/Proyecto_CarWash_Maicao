import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("Falta la librería 'python-docx'. Instala con: pip install python-docx")


def add_title(document: Document, text: str):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_subtitle(document: Document, text: str):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str):
    document.add_paragraph(text)


def add_bullet(document: Document, text: str):
    document.add_paragraph(text, style='List Bullet')


def build_manual(doc: Document):
    # Portada
    add_title(doc, "Manual Técnico - Autolavados Plataforma")
    add_subtitle(doc, f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Entorno: Producción (PythonAnywhere)")

    # Helper para leer variables de entorno desde .env/.env.pythonanywhere si existen
    def load_env_values():
        base_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.dirname(base_dir)
        env_candidates = [
            os.path.join(project_root, '.env'),
            os.path.join(project_root, '.env.pythonanywhere')
        ]
        env = {}
        for path in env_candidates:
            if os.path.exists(path):
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        for line in f:
                            line = line.strip()
                            if not line or line.startswith('#'):
                                continue
                            if '=' in line:
                                k, v = line.split('=', 1)
                                env[k.strip()] = v.strip()
                except Exception:
                    pass
                break
        return env

    env = load_env_values()
    site_url = env.get('SITE_URL', 'https://<usuario>.pythonanywhere.com')
    pa_domain = env.get('PYTHONANYWHERE_DOMAIN', '<usuario>.pythonanywhere.com')
    allowed_hosts = env.get('ALLOWED_HOSTS', '<usuario>.pythonanywhere.com,localhost,127.0.0.1')
    custom_domain = env.get('CUSTOM_DOMAIN', '')

    # Resumen
    add_heading(doc, "1. Resumen del Sistema", level=1)
    add_paragraph(doc, (
        "Sistema web para la gestión integral de un autolavado: autenticación de usuarios, gestión de clientes y vehículos, "
        "reservas de servicios, notificaciones, paneles de control y administración. Basado en Django 4.2 y MySQL, expone APIs REST "
        "para consumo desde frontend y soporta flujos internos mediante vistas clásicas."
    ))

    # Arquitectura
    add_heading(doc, "2. Arquitectura y Tecnologías", level=1)
    add_bullet(doc, "Framework: Django 4.2.x")
    add_bullet(doc, "Base de datos: MySQL (Producción) / SQLite (Desarrollo)")
    add_bullet(doc, "API: Django REST Framework (autenticación, permisos, paginación)")
    add_bullet(doc, "Tiempo real / eventos: Channels (configurado en el proyecto)")
    add_bullet(doc, "Servidor de aplicación: Gunicorn (producción)")
    add_bullet(doc, "Archivos estáticos: WhiteNoise (compresión y manifest)")
    add_bullet(doc, "Almacenamiento de medios: Local por defecto; S3 opcional con django-storages/boto3")
    add_bullet(doc, "Cache: LocMem por defecto; Redis opcional si REDIS_URL está definido")
    add_bullet(doc, "Dependencias clave: Django, asgiref, channels, DRF, cors-headers, mysqlclient, pillow, requests, sqlparse")

    # Estructura de Apps
    add_heading(doc, "3. Estructura de Aplicaciones", level=1)
    add_bullet(doc, "autenticacion: Modelo de usuario personalizado, formularios y vistas de login/logout")
    add_bullet(doc, "clientes: Administración de clientes y vehículos")
    add_bullet(doc, "reservas: Core del sistema (Reservas, Servicios, Bahías, Horarios)")
    add_bullet(doc, "notificaciones: Envío y gestión de notificaciones a clientes y personal")
    add_bullet(doc, "empleados: Gestión de personal, incentivos, reportes y utilidades")
    add_bullet(doc, "dashboards: Vistas públicas y gerenciales para métricas e indicadores")

    # Estructura completa del proyecto (árbol)
    def generate_tree_lines(root_dir: str, max_depth: int = 3, max_entries_per_dir: int = 200):
        lines = []
        def walk(dir_path: str, depth: int, prefix: str = ""):
            try:
                entries = sorted(os.listdir(dir_path))
            except Exception:
                return
            count = 0
            total = len(entries)
            for i, name in enumerate(entries):
                full = os.path.join(dir_path, name)
                is_dir = os.path.isdir(full)
                connector = "└── " if i == total - 1 else "├── "
                line = f"{prefix}{connector}{name}{'/' if is_dir else ''}"
                lines.append(line)
                count += 1
                if count >= max_entries_per_dir:
                    remaining = total - count
                    if remaining > 0:
                        lines.append(f"{prefix}└── ... ({remaining} más)")
                    break
                if is_dir and depth < max_depth:
                    new_prefix = f"{prefix}{'    ' if i == total - 1 else '│   '}"
                    walk(full, depth + 1, new_prefix)
        lines.append(os.path.basename(root_dir) + "/")
        walk(root_dir, 1, "")
        return lines

    add_heading(doc, "3.1 Estructura completa del proyecto", level=2)
    project_root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    for line in generate_tree_lines(project_root_dir, max_depth=3):
        add_paragraph(doc, line)

    # Modelos principales
    add_heading(doc, "4. Modelos principales (visión general)", level=1)
    add_bullet(doc, "Usuario (autenticacion.Usuario): Extiende el modelo de Django para roles y atributos del sistema")
    add_bullet(doc, "Cliente y Vehiculo (clientes): Relación y administración de la flota de clientes")
    add_bullet(doc, "Reserva (reservas): Entidad clave con estado, servicio, medio de pago y asignaciones")
    add_bullet(doc, "Servicio y Bahia (reservas): Catálogo de servicios y espacios físicos de atención")
    add_bullet(doc, "Notificacion (notificaciones): Mensajes informativos, estados y tipos")

    # Middleware
    add_heading(doc, "5. Middleware personalizado", level=1)
    add_bullet(doc, "CSRFDebugMiddleware: Diagnóstico de tokens CSRF en flujos AJAX")
    add_bullet(doc, "AJAXExceptionMiddleware: Manejo centralizado de excepciones para solicitudes AJAX")
    add_bullet(doc, "LoginRequiredMiddleware: Redirección a login para vistas protegidas")
    add_bullet(doc, "TimezoneMiddleware: Gestión de zona horaria del usuario")

    # Configuración de producción
    add_heading(doc, "6. Configuración de Producción (PythonAnywhere)", level=1)
    add_bullet(doc, "Archivo: settings_production.py")
    add_bullet(doc, f"Dominio PythonAnywhere: {pa_domain}")
    if custom_domain:
        add_bullet(doc, f"Dominio personalizado: {custom_domain}")
    add_bullet(doc, f"SITE_URL: {site_url}")
    add_bullet(doc, f"ALLOWED_HOSTS: {allowed_hosts}")
    add_bullet(doc, "Variables de entorno: SECRET_KEY, DB_NAME/USER/PASSWORD/HOST/PORT, EMAIL_*, DEBUG=false")
    add_bullet(doc, "Estáticos: STATIC_ROOT y WhiteNoise para servir assets")
    add_bullet(doc, "Email: SMTP por defecto (configurable)")

    # API y URLs
    add_heading(doc, "7. API y Ruteo", level=1)
    add_bullet(doc, "URLs principales en autolavados_plataforma.urls")
    add_bullet(doc, "APIs REST: vistas y endpoints en apps (ej. reservas/urls_api.py, autenticacion/urls.py)")
    add_bullet(doc, "Autenticación API: Token, Basic y Session (REST_FRAMEWORK)")

    # Endpoints clave
    add_heading(doc, "7.1 Endpoints de Autenticación", level=2)
    add_bullet(doc, f"{site_url}/api/auth/api/registro/ [POST] - Registro de nuevos usuarios")
    add_bullet(doc, f"{site_url}/api/auth/api/login/ [POST] - Inicio de sesión")
    add_bullet(doc, f"{site_url}/api/auth/api/logout/ [POST] - Cierre de sesión")
    add_bullet(doc, f"{site_url}/api/auth/api/perfil/ [GET/PUT] - Ver/actualizar perfil")

    add_heading(doc, "7.2 Endpoints de Reservas", level=2)
    add_bullet(doc, f"{site_url}/api/reservas/servicios/ [GET/POST] - CRUD de servicios")
    add_bullet(doc, f"{site_url}/api/reservas/reservas/ [GET/POST] - CRUD de reservas")
    add_bullet(doc, f"{site_url}/api/reservas/bahias/ [GET] - Listado de bahías")
    add_bullet(doc, f"{site_url}/api/reservas/verificar-placa/ [POST] - Validar placa única")

    add_heading(doc, "7.3 Endpoints de Notificaciones", level=2)
    add_bullet(doc, f"{site_url}/notificaciones/api/contador/ [GET] - Conteo de notificaciones")
    add_bullet(doc, f"{site_url}/notificaciones/api/dropdown/ [GET] - Dropdown de notificaciones")
    add_bullet(doc, f"{site_url}/notificaciones/api/marcar-leida/<id>/ [POST] - Marcar leída")
    add_bullet(doc, f"{site_url}/notificaciones/api/marcar-todas-leidas/ [POST] - Marcar todas leídas")

    # Dependencias
    add_heading(doc, "8. Dependencias y utilidades", level=1)
    add_bullet(doc, "requirements.txt: conjunto mínimo para producción")
    add_bullet(doc, "Utilidades: qrcode (códigos QR), reportlab (PDF), openpyxl/xlsxwriter (Excel)")
    add_bullet(doc, "Opcional S3: django-storages + boto3, si se usa almacenamiento en AWS")

    # Migraciones
    add_heading(doc, "9. Migraciones y Base de Datos", level=1)
    add_paragraph(doc, (
        "Las carpetas de migraciones están incluidas en el repositorio para garantizar que el entorno de producción pueda "
        "aplicar los cambios de esquema correctamente. Flujo recomendado: makemigrations (desarrollo), commit/push, migrate "
        "en producción."
    ))

    # Seguridad
    add_heading(doc, "10. Seguridad", level=1)
    add_bullet(doc, "SECRET_KEY se gestiona por variables de entorno y no debe versionarse")
    add_bullet(doc, "DEBUG debe estar en False en producción")
    add_bullet(doc, "CSRF y autenticación: protección para formularios y APIs, sesiones seguras")

    # Pruebas
    add_heading(doc, "11. Pruebas y utilidades de verificación", level=1)
    add_bullet(doc, "Suite de tests localizada en archivos test_*.py (API, reservas, empleados, UI, despliegue)")
    add_bullet(doc, "Scripts de verificación: check_reservas.py, check_notifications.py, verificación de conexiones y entorno")

    # Despliegue
    add_heading(doc, "12. Despliegue en PythonAnywhere (resumen)", level=1)
    add_bullet(doc, "Actualizar código: git pull origin main")
    add_bullet(doc, "Instalar dependencias: pip install --no-cache-dir -r requirements.txt")
    add_bullet(doc, "Migraciones: python manage.py migrate --noinput")
    add_bullet(doc, "Estáticos: python manage.py collectstatic --noinput")
    add_bullet(doc, "Recargar la web app desde el panel de PythonAnywhere")

    # Mantenimiento
    add_heading(doc, "13. Mantenimiento y buenas prácticas", level=1)
    add_bullet(doc, "Separar dependencias de desarrollo si es necesario (requirements-dev.txt)")
    add_bullet(doc, "Control de cuotas: evitar paquetes pesados innecesarios en producción")
    add_bullet(doc, "Revisar ALLOWED_HOSTS y variables de entorno en cada despliegue")

    # Diagramas
    add_heading(doc, "14. Diagramas de Arquitectura y Flujo", level=1)
    add_bullet(doc, "Arquitectura: ver docs/diagrams/arquitectura.svg")
    add_bullet(doc, "Flujo de reservas: ver docs/diagrams/flujo_reservas.svg")


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(base_dir)
    docs_dir = os.path.join(project_root, 'docs')
    os.makedirs(docs_dir, exist_ok=True)

    output_path = os.path.join(docs_dir, 'Manual_Tecnico_Autolavados.docx')
    document = Document()
    build_manual(document)
    document.save(output_path)
    print(f"Manual técnico generado en: {output_path}")


if __name__ == '__main__':
    main()